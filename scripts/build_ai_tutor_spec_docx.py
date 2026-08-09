#!/usr/bin/env python3
"""Build the polished AI Tutor comprehensive specification from Markdown."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = RGBColor(16, 24, 40)
PURPLE = RGBColor(108, 76, 241)
CORAL = RGBColor(255, 107, 94)
MUTED = RGBColor(102, 112, 133)
LIGHT_PURPLE = "F1EEFF"
LIGHT_GRAY = "F2F4F7"
GRID = "D0D5DD"
WHITE = RGBColor(255, 255, 255)

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 90, "bottom": 90, "start": 120, "end": 120}


def clean_text(text):
    replacements = {
        "—": "-", "–": "-", "→": "->", "≥": ">=", "≤": "<=",
        "×": "x", "“": '"', "”": '"', "’": "'", "…": "...", "·": "|",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return text


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in CELL_MARGIN.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), GRID)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    # Word exposes the first-row header marker to assistive technology. Apply it
    # consistently to both data tables and the compact one-row callout tables
    # used by this document template.
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if header_tr_pr.find(qn("w:tblHeader")) is None:
        table_header = OxmlElement("w:tblHeader")
        table_header.set(qn("w:val"), "true")
        header_tr_pr.append(table_header)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def widths_for_table(column_count):
    if column_count == 2:
        return [2700, 6660]
    if column_count == 3:
        return [2160, 3600, 3600]
    if column_count == 4:
        return [1800, 2400, 2400, 2760]
    base = CONTENT_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run_font(run, size=8.5, color=MUTED)


def new_numbering_id(document, start=1):
    numbering = document.part.numbering_part.element
    base_num_id = document.styles["List Number"]._element.pPr.numPr.numId.val
    base_num = numbering.xpath(f"./w:num[@w:numId='{base_num_id}']")[0]
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    heading_tokens = {
        "Heading 1": (16, PURPLE, 18, 10),
        "Heading 2": (13, PURPLE, 14, 7),
        "Heading 3": (11.5, INK, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.font.color.rgb = INK
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2


def configure_page(document):
    for section in document.sections:
        section.top_margin = Inches(0.78)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("CLAT PREP STUDIO  |  AI TUTOR SPECIFICATION")
        set_run_font(run, size=8, color=MUTED, bold=True)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        label = p.add_run("CONFIDENTIAL WORKING SPEC  |  ")
        set_run_font(label, size=8, color=MUTED)
        add_page_number(p)


def add_inline_markdown(paragraph, text, default_size=10.5, default_color=INK):
    text = clean_text(text)
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=default_size, color=default_color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Courier New", size=9, color=PURPLE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=default_size, color=default_color)


def add_cover(document):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(46)
    p.paragraph_format.space_after = Pt(8)
    kicker = p.add_run("CLAT PREP STUDIO  /  PRODUCT & ENGINEERING")
    set_run_font(kicker, size=9, color=CORAL, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    title = p.add_run("AI Tutor")
    set_run_font(title, size=34, color=INK, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    subtitle = p.add_run("Comprehensive Product, Journey, BYOK and Delivery Specification")
    set_run_font(subtitle, size=17, color=PURPLE, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    lead = p.add_run(
        "A build-ready operating specification for an evidence-led tutor that improves accuracy, "
        "speed, difficulty transfer and score reliability while allowing students to power "
        "conversational coaching with their own model credential."
    )
    set_run_font(lead, size=11.5, color=MUTED)

    table = document.add_table(rows=5, cols=2)
    metadata = [
        ("Version", "1.0 - Build-ready specification"),
        ("Date", "2 August 2026"),
        ("Exam target", "CLAT 2027 - 6 December 2026"),
        ("Outcome", "Maximise the controllable probability of 110+/120"),
        ("Audience", "Product, learning science, design, engineering, privacy and academics"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_PURPLE)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=9.5, color=PURPLE, bold=True)
        for run in row.cells[1].paragraphs[0].runs:
            set_run_font(run, size=9.5, color=INK)
    set_table_geometry(table, [2100, 7260])

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("CORE DECISION")
    set_run_font(run, size=8.5, color=CORAL, bold=True)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(
        "CLAT Prep Studio owns the educational decision engine. The student's BYOK connection "
        "powers optional conversational intelligence."
    )
    set_run_font(run, size=12, color=INK, bold=True)

    document.add_page_break()


def add_contents(document, headings):
    p = document.add_paragraph(style="Heading 1")
    p.add_run("Contents")
    contents_num_id = new_numbering_id(document)
    for heading in headings:
        p = document.add_paragraph(style="List Number")
        apply_numbering(p, contents_num_id)
        text = re.sub(r"^\d+\.\s*", "", heading)
        add_inline_markdown(p, text, default_size=9.5)
    document.add_page_break()


def add_code_block(document, lines):
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    cell.text = ""
    for index, line in enumerate(lines):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(clean_text(line) if line else " ")
        set_run_font(run, name="Courier New", size=8.2, color=INK)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_markdown_table(document, rows):
    cleaned = []
    for row in rows:
        cells = [clean_text(cell.strip()) for cell in row.strip().strip("|").split("|")]
        cleaned.append(cells)
    if len(cleaned) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cleaned[1]):
        del cleaned[1]
    if not cleaned:
        return
    column_count = max(len(row) for row in cleaned)
    table = document.add_table(rows=len(cleaned), cols=column_count)
    for row_index, values in enumerate(cleaned):
        for col_index in range(column_count):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            value = values[col_index] if col_index < len(values) else ""
            add_inline_markdown(
                p,
                value,
                default_size=8.6 if column_count >= 4 else 9.0,
                default_color=WHITE if row_index == 0 else INK,
            )
            if row_index == 0:
                set_cell_shading(cell, "6C4CF1")
                for run in p.runs:
                    run.bold = True
            elif row_index % 2 == 0:
                set_cell_shading(cell, "FAFAFF")
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    table_header = header_tr_pr.find(qn("w:tblHeader"))
    if table_header is None:
        table_header = OxmlElement("w:tblHeader")
        table_header.set(qn("w:val"), "true")
        header_tr_pr.append(table_header)
    set_table_geometry(table, widths_for_table(column_count))
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_body(document, markdown):
    lines = markdown.splitlines()
    start = 0
    for index, line in enumerate(lines):
        if line.strip() == "---":
            start = index + 1
            break
    lines = lines[start:]
    page_break_sections = {
        "## 5. End-to-end student progress journey",
        "## 10. Conversational intelligence specification",
        "## 14. Privacy, minors and trust controls",
        "## Appendix A — Recommended tutor system-policy outline",
    }

    index = 0
    current_numbering_id = None
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped or stripped == "---":
            current_numbering_id = None
            index += 1
            continue
        if line in page_break_sections:
            document.add_page_break()
        if stripped.startswith("```"):
            current_numbering_id = None
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(document, code_lines)
            index += 1
            continue
        if stripped.startswith("|"):
            current_numbering_id = None
            table_rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(lines[index])
                index += 1
            add_markdown_table(document, table_rows)
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            current_numbering_id = None
            level = min(len(heading_match.group(1)) - 1, 3)
            p = document.add_paragraph(style=f"Heading {level}")
            add_inline_markdown(
                p,
                heading_match.group(2),
                default_size={1: 16, 2: 13, 3: 11.5}[level],
                default_color=PURPLE if level < 3 else INK,
            )
            for run in p.runs:
                run.bold = True
            index += 1
            continue
        if stripped.startswith(">"):
            current_numbering_id = None
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            table = document.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            set_cell_shading(cell, LIGHT_PURPLE)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_markdown(p, " ".join(quote_lines), default_size=10.5, default_color=PURPLE)
            for run in p.runs:
                run.italic = True
            set_table_geometry(table, [CONTENT_WIDTH_DXA])
            index += 0
            continue
        if re.match(r"^-\s+", stripped):
            current_numbering_id = None
            p = document.add_paragraph(style="List Bullet")
            add_inline_markdown(p, re.sub(r"^-\s+", "", stripped))
            index += 1
            continue
        numbered_match = re.match(r"^(\d+)\.\s+", stripped)
        if numbered_match:
            if current_numbering_id is None:
                current_numbering_id = new_numbering_id(document, start=int(numbered_match.group(1)))
            p = document.add_paragraph(style="List Number")
            apply_numbering(p, current_numbering_id)
            add_inline_markdown(p, re.sub(r"^\d+\.\s+", "", stripped))
            index += 1
            continue

        current_numbering_id = None
        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            nxt = lines[index].strip()
            if not nxt or nxt.startswith(("#", "```", "|", ">", "- ")) or re.match(r"^\d+\.\s+", nxt):
                break
            paragraph_lines.append(nxt)
            index += 1
        p = document.add_paragraph()
        add_inline_markdown(p, " ".join(paragraph_lines))


def build(markdown_path, output_path):
    markdown = markdown_path.read_text(encoding="utf-8")
    document = Document()
    configure_styles(document)
    configure_page(document)
    add_cover(document)

    top_headings = [
        clean_text(match.group(1))
        for match in re.finditer(r"^## (\d+\. .+)$", markdown, flags=re.MULTILINE)
    ]
    add_contents(document, top_headings)
    parse_body(document, markdown)

    core = document.core_properties
    core.title = "CLAT Prep Studio AI Tutor - Comprehensive Product Specification"
    core.subject = "AI Tutor journey, BYOK, adaptive learning and delivery specification"
    core.author = "CLAT Prep Studio"
    core.keywords = "CLAT, AI tutor, adaptive learning, BYOK, product specification"
    core.comments = "Generated from the canonical Markdown specification."

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_ai_tutor_spec_docx.py INPUT.md OUTPUT.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
