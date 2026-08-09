#!/usr/bin/env python3
"""Build the polished Perfect Mock Pipeline specification from Markdown."""

import importlib.util
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
BASE_BUILDER_PATH = REPO_ROOT / "scripts" / "build_ai_tutor_spec_docx.py"
spec = importlib.util.spec_from_file_location("base_docx_builder", BASE_BUILDER_PATH)
base = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(base)


# compact_reference_guide preset tokens
INK = RGBColor(11, 37, 69)
BODY = RGBColor(31, 41, 55)
HEADING = RGBColor(46, 116, 181)
HEADING_DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(102, 112, 133)
CAUTION = RGBColor(122, 90, 0)
RISK = RGBColor(155, 28, 28)
TABLE_HEADER = "E8EEF5"
TABLE_ALT = "F8FAFC"
CALLOUT = "F4F6F9"
GRID = "CBD5E1"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), name)
    r_pr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def patch_numbering_level(document, style_name, marker_kind):
    """Apply the preset's real list alignment to the style numbering definition."""
    style = document.styles[style_name]
    num_pr = style._element.pPr.numPr
    if num_pr is None:
        return
    num_id = num_pr.numId.val
    numbering = document.part.numbering_part.element
    num = numbering.xpath(f"./w:num[@w:numId='{num_id}']")
    if not num:
        return
    abstract_id = num[0].find(qn("w:abstractNumId")).get(qn("w:val"))
    abstract = numbering.xpath(f"./w:abstractNum[@w:abstractNumId='{abstract_id}']")
    if not abstract:
        return
    level = abstract[0].find(qn("w:lvl"))
    if level is None:
        return
    p_pr = level.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        level.append(p_pr)
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    for node in list(tabs):
        tabs.remove(node)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = p_pr.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        p_pr.append(indent)
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    if marker_kind == "bullet":
        num_fmt = level.find(qn("w:numFmt"))
        lvl_text = level.find(qn("w:lvlText"))
        if num_fmt is not None:
            num_fmt.set(qn("w:val"), "bullet")
        if lvl_text is not None:
            lvl_text.set(qn("w:val"), "•")


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, HEADING, 18, 10),
        "Heading 2": (13, HEADING, 14, 7),
        "Heading 3": (12, HEADING_DARK, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = BODY
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    patch_numbering_level(document, "List Bullet", "bullet")
    patch_numbering_level(document, "List Number", "decimal")


def configure_page(document):
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("CLAT PREP STUDIO  |  PERFECT MOCK PIPELINE SPECIFICATION")
        set_run_font(run, size=8, color=MUTED, bold=True)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        label = p.add_run("BASELINE DESIGN AUTHORITY  |  ")
        set_run_font(label, size=8, color=MUTED)
        base.add_page_number(p)


def add_cover(document):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run("ENGINEERING SPECIFICATION"), size=9, color=CAUTION, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("Perfect Mock-Paper Pipeline"), size=26, color=INK, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_run_font(
        p.add_run("Digitisation, evidence, review, publication and adaptive-readiness for CLAT Prep Studio"),
        size=14,
        color=HEADING,
        bold=True,
    )

    metadata = [
        ("Version", "1.0 - Baseline design authority"),
        ("Date", "3 August 2026"),
        ("Primary objective", "Build a verified, provenance-complete CLAT question layer"),
        ("Current baseline", "202 PDFs | 8,925 pages | 11,138 candidates | 492 verified items"),
        ("Boundary", "Canonical content release; learner-product stitching follows separately"),
        ("Audience", "Academic, content operations, product, engineering, learning science and governance"),
    ]
    table = document.add_table(rows=len(metadata), cols=2)
    for row, (label, value) in zip(table.rows, metadata):
        row.cells[0].text = label
        row.cells[1].text = value
        base.set_cell_shading(row.cells[0], TABLE_HEADER)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=9.2, color=HEADING_DARK, bold=True)
        for run in row.cells[1].paragraphs[0].runs:
            set_run_font(run, size=9.2, color=BODY)
    base.set_table_geometry(table, [2700, 6660])

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run("DESIGN AUTHORITY"), size=8.5, color=CAUTION, bold=True)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run_font(
        p.add_run(
            "No OCR engine or semantic model may independently publish a scored question. "
            "Source evidence, deterministic gates and accountable review decide learner readiness."
        ),
        size=12,
        color=INK,
        bold=True,
    )
    document.add_page_break()


def add_inline_markdown(paragraph, text, default_size=11, default_color=BODY):
    text = base.clean_text(text)
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=default_size, color=default_color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Courier New", size=9, color=HEADING_DARK)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=default_size, color=default_color)


def widths_for_table(column_count):
    if column_count == 2:
        return [2700, 6660]
    if column_count == 3:
        return [2160, 3600, 3600]
    if column_count == 4:
        return [1800, 2460, 2460, 2640]
    width = CONTENT_WIDTH_DXA // column_count
    widths = [width] * column_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(document, rows):
    cleaned = [[base.clean_text(cell.strip()) for cell in row.strip().strip("|").split("|")] for row in rows]
    if len(cleaned) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cleaned[1]):
        del cleaned[1]
    if not cleaned:
        return
    column_count = max(len(row) for row in cleaned)
    table = document.add_table(rows=len(cleaned), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row_index, values in enumerate(cleaned):
        for col_index in range(column_count):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            value = values[col_index] if col_index < len(values) else ""
            add_inline_markdown(p, value, default_size=8.4 if column_count >= 4 else 8.9, default_color=BODY)
            if row_index == 0:
                base.set_cell_shading(cell, TABLE_HEADER)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = HEADING_DARK
            elif row_index % 2 == 0:
                base.set_cell_shading(cell, TABLE_ALT)
    base.set_table_geometry(table, widths_for_table(column_count))
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def set_table_no_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def add_contents(document, headings):
    p = document.add_paragraph(style="Heading 1")
    p.add_run("Contents")
    midpoint = (len(headings) + 1) // 2
    left = headings[:midpoint]
    right = headings[midpoint:]
    table = document.add_table(rows=midpoint, cols=2)
    for row_index in range(midpoint):
        for col_index, group in enumerate((left, right)):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.05
            if row_index < len(group):
                add_inline_markdown(p, group[row_index], default_size=8.5, default_color=BODY)
    base.set_table_geometry(table, [4680, 4680])
    set_table_no_borders(table)
    document.add_page_break()


def add_manual_decimal(document, stripped):
    """Render explicit source numbering reliably across Word and LibreOffice."""
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_inline_markdown(p, stripped)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(2)
    run = spacer.add_run("\u200b")
    set_run_font(run, size=1, color=BODY)


def parse_body(document, markdown):
    lines = markdown.splitlines()
    start = 0
    for index, line in enumerate(lines):
        if line.strip() == "---":
            start = index + 1
            break
    lines = lines[start:]

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped or stripped == "---":
            index += 1
            continue
        if stripped.startswith("```"):
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            base.add_code_block(document, code_lines)
            index += 1
            continue
        if stripped.startswith("|"):
            table_rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(lines[index])
                index += 1
            add_markdown_table(document, table_rows)
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)) - 1, 3)
            p = document.add_paragraph(style=f"Heading {level}")
            add_inline_markdown(
                p,
                heading_match.group(2),
                default_size={1: 16, 2: 13, 3: 12}[level],
                default_color=HEADING if level < 3 else HEADING_DARK,
            )
            for run in p.runs:
                run.bold = True
            index += 1
            continue
        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            table = document.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            base.set_cell_shading(cell, CALLOUT)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_markdown(p, " ".join(quote_lines), default_size=10.5, default_color=HEADING)
            for run in p.runs:
                run.italic = True
            base.set_table_geometry(table, [CONTENT_WIDTH_DXA])
            continue
        if re.match(r"^-\s+", stripped):
            p = document.add_paragraph(style="List Bullet")
            add_inline_markdown(p, re.sub(r"^-\s+", "", stripped))
            index += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            add_manual_decimal(document, stripped)
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            nxt = lines[index].strip()
            if not nxt or nxt.startswith(("#", "```", "|", ">", "- ")) or re.match(r"^\d+\.\s+", nxt):
                break
            paragraph_lines.append(nxt)
            index += 1
        p = document.add_paragraph()
        add_inline_markdown(p, " ".join(paragraph_lines))


def build(markdown_path, output_path):
    markdown = markdown_path.read_text(encoding="utf-8")
    document = Document()

    # Rebind the generic parser's formatting hooks to this preset.
    base.INK = BODY
    base.PURPLE = HEADING
    base.MUTED = MUTED
    base.LIGHT_PURPLE = CALLOUT
    base.LIGHT_GRAY = "F2F4F7"
    base.GRID = GRID
    base.CONTENT_WIDTH_DXA = CONTENT_WIDTH_DXA
    base.TABLE_INDENT_DXA = TABLE_INDENT_DXA
    base.set_run_font = set_run_font
    base.add_inline_markdown = add_inline_markdown
    base.add_markdown_table = add_markdown_table
    base.widths_for_table = widths_for_table

    configure_styles(document)
    configure_page(document)
    add_cover(document)

    headings = [
        base.clean_text(match.group(1))
        for match in re.finditer(r"^## (\d+\. .+)$", markdown, flags=re.MULTILINE)
    ]
    add_contents(document, headings)
    parse_body(document, markdown)

    core = document.core_properties
    core.title = "CLAT Prep Studio - Perfect Mock-Paper Pipeline Engineering Specification"
    core.subject = "Digitisation, OCR, classification, review, publication and adaptive-readiness"
    core.author = "CLAT Prep Studio"
    core.keywords = "CLAT, PDF digitisation, OCR, question bank, adaptive learning, pipeline"
    core.comments = "Generated from the canonical Markdown design authority."

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_perfect_mock_pipeline_spec_docx.py INPUT.md OUTPUT.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
